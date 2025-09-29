"""
Procesador de archivos de contexto para extraer contenido de diferentes formatos
"""
import json
import os
import logging
from typing import Dict, Any, Optional, Tuple
from datetime import datetime, timedelta
from pathlib import Path

import docx
from django.conf import settings
from django.core.files.storage import default_storage
from langchain_openai import ChatOpenAI

logger = logging.getLogger(__name__)


class ContextFileProcessor:
    """
    Procesador de archivos de contexto que extrae contenido de diferentes formatos
    """

    def __init__(self):
        self.llm = ChatOpenAI(
            model="gpt-4o-mini",
            temperature=0.1,
            openai_api_key=settings.OPENAI_API_KEY
        ) if hasattr(settings, 'OPENAI_API_KEY') else None

    def process_file(self, context_file) -> Tuple[bool, str, Dict[str, Any]]:
        """
        Procesa un archivo de contexto según su tipo

        Args:
            context_file: Instancia de ContextFile

        Returns:
            Tuple[bool, str, Dict]: (success, extracted_content, metadata)
        """
        try:
            logger.info(f"Procesando archivo: {context_file.name} ({context_file.file_type})")

            # Obtener ruta del archivo
            file_path = context_file.file.path

            # Procesar según tipo
            if context_file.file_type == 'json':
                return self._process_json(file_path)
            elif context_file.file_type == 'txt':
                return self._process_txt(file_path)
            elif context_file.file_type == 'docx':
                return self._process_docx(file_path)
            elif context_file.file_type == 'pdf':
                return self._process_pdf(file_path)
            else:
                return False, "Tipo de archivo no soportado", {}

        except Exception as e:
            logger.error(f"Error procesando archivo {context_file.name}: {e}")
            return False, f"Error procesando archivo: {str(e)}", {}

    def _process_json(self, file_path: str) -> Tuple[bool, str, Dict[str, Any]]:
        """Procesa archivos JSON"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # Convertir JSON a texto estructurado
            content = self._json_to_text(data)

            metadata = {
                'keys_count': len(data) if isinstance(data, dict) else None,
                'items_count': len(data) if isinstance(data, list) else None,
                'structure_type': 'object' if isinstance(data, dict) else 'array' if isinstance(data, list) else 'value'
            }

            return True, content, metadata

        except json.JSONDecodeError as e:
            return False, f"Error decodificando JSON: {str(e)}", {}
        except Exception as e:
            return False, f"Error procesando JSON: {str(e)}", {}

    def _process_txt(self, file_path: str) -> Tuple[bool, str, Dict[str, Any]]:
        """Procesa archivos de texto"""
        try:
            # Intentar diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            content = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                return False, "No se pudo decodificar el archivo de texto", {}

            # Calcular metadatos
            lines = content.split('\n')
            words = content.split()

            metadata = {
                'lines_count': len(lines),
                'words_count': len(words),
                'characters_count': len(content),
                'encoding_used': encoding
            }

            return True, content, metadata

        except Exception as e:
            return False, f"Error procesando TXT: {str(e)}", {}

    def _process_docx(self, file_path: str) -> Tuple[bool, str, Dict[str, Any]]:
        """Procesa archivos Word"""
        try:
            doc = docx.Document(file_path)

            # Extraer texto de párrafos
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            content = '\n\n'.join(paragraphs)

            # Extraer texto de tablas
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(' | '.join(row_data))
                if table_data:
                    tables_content.append('\n'.join(table_data))

            if tables_content:
                content += '\n\n--- TABLAS ---\n\n' + '\n\n'.join(tables_content)

            metadata = {
                'paragraphs_count': len(paragraphs),
                'tables_count': len(doc.tables),
                'words_count': len(content.split()),
                'characters_count': len(content)
            }

            return True, content, metadata

        except Exception as e:
            return False, f"Error procesando DOCX: {str(e)}", {}

    def _process_pdf(self, file_path: str) -> Tuple[bool, str, Dict[str, Any]]:
        """Procesa archivos PDF"""
        try:
            import PyPDF2

            content_pages = []

            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            content_pages.append(f"--- Página {page_num + 1} ---\n{text}")
                    except Exception as e:
                        logger.warning(f"Error extrayendo página {page_num + 1}: {e}")
                        continue

            content = '\n\n'.join(content_pages)

            metadata = {
                'pages_count': len(pdf_reader.pages),
                'pages_with_text': len(content_pages),
                'words_count': len(content.split()),
                'characters_count': len(content)
            }

            return True, content, metadata

        except ImportError:
            return False, "PyPDF2 no está instalado. Instalar con: pip install PyPDF2", {}
        except Exception as e:
            return False, f"Error procesando PDF: {str(e)}", {}

    def _json_to_text(self, data, level=0) -> str:
        """Convierte estructura JSON a texto legible"""
        indent = "  " * level

        if isinstance(data, dict):
            lines = []
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    lines.append(f"{indent}{key}:")
                    lines.append(self._json_to_text(value, level + 1))
                else:
                    lines.append(f"{indent}{key}: {value}")
            return '\n'.join(lines)

        elif isinstance(data, list):
            lines = []
            for i, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    lines.append(f"{indent}[{i}]:")
                    lines.append(self._json_to_text(item, level + 1))
                else:
                    lines.append(f"{indent}[{i}]: {item}")
            return '\n'.join(lines)

        else:
            return f"{indent}{data}"

    def generate_summary(self, content: str, max_length: int = 500) -> str:
        """
        Genera un resumen del contenido usando LLM
        """
        if not self.llm:
            # Si no hay LLM, generar resumen simple
            words = content.split()
            if len(words) <= 50:
                return content
            return ' '.join(words[:50]) + "..."

        try:
            prompt = f"""
            Genera un resumen conciso del siguiente contenido en español, máximo {max_length} caracteres:

            {content[:2000]}  # Limitar entrada para evitar tokens excesivos

            Resumen:
            """

            response = self.llm.invoke([{"role": "user", "content": prompt}])
            summary = response.content.strip()

            # Truncar si es necesario
            if len(summary) > max_length:
                summary = summary[:max_length - 3] + "..."

            return summary

        except Exception as e:
            logger.error(f"Error generando resumen con LLM: {e}")
            # Fallback a resumen simple
            words = content.split()
            if len(words) <= 50:
                return content
            return ' '.join(words[:50]) + "..."

    def validate_file(self, file_path: str, file_type: str) -> Tuple[bool, str]:
        """
        Valida que el archivo sea del tipo esperado y esté accesible
        """
        try:
            if not os.path.exists(file_path):
                return False, "El archivo no existe"

            if not os.path.isfile(file_path):
                return False, "La ruta no apunta a un archivo"

            file_size = os.path.getsize(file_path)
            max_size = 50 * 1024 * 1024  # 50MB máximo

            if file_size > max_size:
                return False, f"Archivo demasiado grande ({file_size / 1024 / 1024:.1f}MB). Máximo 50MB"

            # Validaciones específicas por tipo
            if file_type == 'json':
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        json.load(f)
                except json.JSONDecodeError:
                    return False, "Archivo JSON no válido"

            return True, "Archivo válido"

        except Exception as e:
            return False, f"Error validando archivo: {str(e)}"


# Función helper para usar desde views/tasks
def process_context_file(context_file_id: int) -> bool:
    """
    Procesa un archivo de contexto por ID
    """
    from apps.chat.models import ContextFile, ContextFileProcessingLog

    try:
        context_file = ContextFile.objects.get(id=context_file_id)
        processor = ContextFileProcessor()

        # Marcar como procesando
        context_file.status = 'processing'
        context_file.save()

        # Log inicio
        log_entry = ContextFileProcessingLog.objects.create(
            context_file=context_file,
            action='extract_content',
            status='started',
            details='Iniciando extracción de contenido'
        )

        start_time = datetime.now()

        # Procesar archivo
        success, content, metadata = processor.process_file(context_file)

        execution_time = datetime.now() - start_time

        if success:
            # Generar resumen
            summary = processor.generate_summary(content)

            # Actualizar modelo
            context_file.extracted_content = content
            context_file.content_summary = summary
            context_file.metadata.update(metadata)
            context_file.status = 'processed'
            context_file.processing_error = ''

            # Log éxito
            log_entry.status = 'completed'
            log_entry.details = f'Contenido extraído exitosamente. {len(content)} caracteres.'
            log_entry.execution_time = execution_time

        else:
            # Error en procesamiento
            context_file.status = 'error'
            context_file.processing_error = content  # content contiene el mensaje de error

            # Log error
            log_entry.status = 'failed'
            log_entry.details = content
            log_entry.execution_time = execution_time

        context_file.save()
        log_entry.save()

        return success

    except Exception as e:
        logger.error(f"Error procesando context_file {context_file_id}: {e}")
        return False